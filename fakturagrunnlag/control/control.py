import logging
import csv
import os

from PyQt5.QtWidgets import QApplication, QMessageBox
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt
from openpyxl import Workbook
from openpyxl.utils import get_column_letter

from reportlab.lib.pagesizes import A4
from reportlab.pdfbase.pdfmetrics import stringWidth
from reportlab.pdfgen import canvas
from reportlab.lib.units import cm

from common.gui.utils import show_message
from fakturagrunnlag.brreg import brreg
from fakturagrunnlag.db import sql
from fakturagrunnlag.eventor import eventor
from common.html_builder import HtmlBuilder


def import_eventor_clubs(parent, eventor_apikey, progress):
    logging.info("import_eventor_clubs")
    orgs = eventor.get_clubs(eventor_apikey, "NOR")
    actual_progress = 0
    progress.setValue(actual_progress)
    sql.delete_inv_customers(parent.ctx.conn_mgr, "NOR")

    progress.setMaximum(len(orgs))
    for i, org in enumerate(orgs, start=1):
        sql.insert_inv_customer(parent.ctx.conn_mgr, org)
        progress.setValue(i)
        QApplication.processEvents()  # holder GUI responsivt

    return len(orgs)

"""
    invoice_date = datetime.date.today()
    due_date = invoice_date + datetime.timedelta(days=30)
    sql.cre_order_bundle(parent.ctx.conn_mgr, invoice_date, due_date, "BO", "BF", "Samlefaktura Veteranmesterskapet.", "NOK")
"""

def export_tripletex_csv(parent, bundle_id):
    logging.info("control.export_tripletex")
    rows, columns = sql.export_tripletex(parent.ctx.conn_mgr, bundle_id, parent.get_order_no_base(), parent.get_customer_no_base())

    download_path = get_download_path()
    full_path = os.path.join(download_path, "tripletex_invoice.csv")

    write_tripletex_csv(rows, columns, full_path)
    return len(rows)

"""
Henter organisasjonsnummer fra Brønnøysund registeret og
plasserer på alle ordrer i en bunt.
"""
def add_org_no(parent, bundle_id, progress):
#    logging.info("control.add_org_no")
    rows, columns = sql.select_norge_orders(parent.ctx.conn_mgr, bundle_id)

    actual_progress = 0
    progress.setValue(actual_progress)
    progress.setMaximum(len(rows))

    i = 0 # Progress
    j = 0 # Antall org_no funnet.
    for row in rows:
        order_id = row[0]
        org_name = row[3]
        org_no = brreg.get_orgnummer(org_name)
        i = i + 1
        if org_no:
            sql.upd_org_no(parent.ctx.conn_mgr, order_id, org_no)
            j = j + 1
        progress.setValue(i)
        QApplication.processEvents()  # holder GUI responsivt

    return j


def write_tripletex_csv(rows, columns, output_path="tripletex_invoice.csv"):
    """
    Writes Tripletex-compatible CSV from SQL query result.

    Parameters:
        rows (list of tuples): Result rows from cursor.fetchall()
        columns (list of str): Column names from cursor.description
        output_path (str): Path to save the CSV file
    """
    # Convert column names: 'invoice_no' → 'INVOICE NO'
    header = [col.replace("_", " ").upper() for col in columns]


    with open(output_path, mode="w", newline="", encoding="utf-8-sig") as file:
        writer = csv.writer(file, delimiter=";")
        writer.writerow(header)
        for row in rows:
            writer.writerow(row)

def export_tripletex_excel(parent, bundle_id):
    logging.info("control.export_tripletex")
    rows, columns = sql.export_tripletex(parent.ctx.conn_mgr, bundle_id, parent.get_order_no_base(), parent.get_customer_no_base())

    download_path = get_download_path()
    full_path = os.path.join(download_path, "tripletex_invoice.xlsx")

    write_tripletex_excel(parent, rows, columns, full_path)
    return len(rows)


def write_tripletex_excel(parent, rows, columns, full_path):
    """
    Eksporterer data til en Excel-fil med kolonneoverskrifter.

    :param rows: Liste av rader (hver rad er en liste med verdier)
    :param columns: Liste med kolonnenavn
    :param full_path: Full filsti til .xlsx-filen
    """

    # Convert column names: 'invoice_no' → 'INVOICE NO'
    header = [col.replace("_", " ").upper() for col in columns]

    wb = Workbook()
    ws = wb.active
    ws.title = "Fakturagrunnlag"

    # Skriv kolonneoverskrifter
    ws.append(header)

    # Skriv rader
    for row in rows:
        ws.append(row)

    # Autojuster kolonnebredder
    for i, col in enumerate(columns, start=1):
        max_len = max((len(str(cell)) for cell in [col] + [r[i-1] for r in rows if len(r) > i-1]), default=0)
        ws.column_dimensions[get_column_letter(i)].width = max(10, min(max_len + 2, 40))

    try:
        wb.save(full_path)
    except PermissionError:
        QMessageBox.warning(
            parent,
            "Feil ved lagring",
            f"Kunne ikke lagre filen:\n{full_path}\n\n"
            "Sannsynligvis er den allerede åpen i Excel.\n"
            "Lukk filen og prøv igjen."
        )


def get_download_path():
    return os.path.join(os.path.expanduser("~"), "Downloads")




def calculate_mod10(digits: str) -> str:
    """
    Calculates MOD10 check digit for a numeric string.
    """
    total = 0
    factors = [2, 1]
    for i, digit in enumerate(reversed(digits)):
        product = int(digit) * factors[i % 2]
        total += sum(int(d) for d in str(product))
    check_digit = (10 - (total % 10)) % 10
    return str(check_digit)

def calculate_mod11(digits: str) -> str:
    """
    Calculates MOD11 check digit for a numeric string.
    Returns '-' if MOD11 is not valid for the input.
    """
    weights = [2, 3, 4, 5, 6, 7]
    total = 0
    for i, digit in enumerate(reversed(digits)):
        total += int(digit) * weights[i % len(weights)]
    remainder = total % 11
    if remainder == 0:
        return '0'
    elif remainder == 1:
        return '-'  # Invalid MOD11 result
    else:
        return str(11 - remainder)

def generate_kid(base: str, method: str = "MOD10") -> str:
    """
    Generates a KID number using the specified check digit method.
    Parameters:
        base (str): The numeric base (e.g. invoice number or customer ID)
        method (str): 'MOD10' or 'MOD11'
    Returns:
        str: KID number with check digit
    Raises:
        ValueError: If input is invalid or MOD11 fails
    """
    base_digits = ''.join(filter(str.isdigit, base))
    if not base_digits:
        raise ValueError("Base must contain at least one digit.")

    if method == "MOD10":
        check = calculate_mod10(base_digits)
    elif method == "MOD11":
        check = calculate_mod11(base_digits)
        if check == '-':
            raise ValueError("MOD11 check digit is invalid for this base.")
    else:
        raise ValueError("Invalid method. Choose 'MOD10' or 'MOD11'.")

    return base_digits + check

def add_race_to_bundle(parent, bundle_id, raceid):
    count_orders = sql.cre_orders(parent.ctx.conn_mgr, bundle_id, raceid)
    count_lines = sql.cre_order_lines(parent.ctx.conn_mgr, bundle_id, raceid)
    sql.append_remark(parent.ctx.conn_mgr, bundle_id, raceid)
    sql.upd_race_bundle_id(parent.ctx.conn_mgr, raceid, bundle_id)
    return str(count_orders), str(count_lines)

def remove_race_from_bundle(parent, bundle_id, raceid):
    count_lines = sql.rem_order_lines(parent.ctx.conn_mgr, bundle_id, raceid)
    count_orders = sql.rem_empty_orders(parent.ctx.conn_mgr, bundle_id)
    sql.remove_race_from_remark(parent.ctx.conn_mgr, bundle_id, raceid)
    sql.upd_race_bundle_id(parent.ctx.conn_mgr, raceid, None)
    return str(count_orders), str(count_lines)

def delete_bundle(parent, bundle_id):
    logging.info("control.delete_bundle")
    sql.delete_order_lines(parent.ctx.conn_mgr, bundle_id)
    sql.delete_orders(parent.ctx.conn_mgr, bundle_id)
    sql.delete_bundle(parent.ctx.conn_mgr, bundle_id)
    sql.remove_bundle_from_race(parent.ctx.conn_mgr, bundle_id)
    parent.load_bundles()
    parent.line_table.setRowCount(0)
    parent.order_table.setRowCount(0)


def make_order_word(parent, invoice_config, order_id):
    logging.info("control.make_order_word")
    order_no_base = invoice_config.getint("ordrenummer_start", fallback=100000)
    rows, columns = sql.select_order(parent.ctx.conn_mgr, order_id, order_no_base)
    download_path = get_download_path()
    write_manual_invoice_word(parent, invoice_config, rows, columns, download_path)

def make_order_pdf(parent, invoice_config, order_id):
    logging.info("control.make_order_pdf")
    order_no_base = invoice_config.getint("ordrenummer_start", fallback=100000)
    rows, columns = sql.select_order(parent.ctx.conn_mgr, order_id, order_no_base)
    download_path = get_download_path()
    write_manual_invoice_pdf(parent, invoice_config, rows, columns, download_path)

def write_manual_invoice_word(parent, invoice_config, rows, columns, download_path):
    doc = Document()

    # Marger
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.left_margin = Cm(1.0)
        section.right_margin = Cm(1.0)
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)

    # Tittel: FAKTURA i font size 24
    title = doc.add_paragraph()
    run = title.add_run("ORDRE")
    run.font.size = Pt(24)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Fellesdata fra første rad (mottaker)
    first = rows[0]
    felt = dict(zip(columns, first))

    # Header-tabell med to kolonner
    header = doc.add_table(rows=1, cols=2)
    header.autofit = False
    hdr_cells = header.rows[0].cells

    # Venstre: mottaker
    left = hdr_cells[0].paragraphs[0]
    left.add_run(f"{felt['Klubb']}\n")
    if felt['Adresse']:
        left.add_run(f"{felt['Adresse']}\n")
    if felt['E_post']:  # ledetekst fjernet, men viser e-post hvis den finnes
        left.add_run(f"{felt['E_post']}\n")
    if felt['Telefonnr']:
        left.add_run(f"Tlf: {felt['Telefonnr']}")

    # Høyre: utsteder + fakturainfo
    right = hdr_cells[1].paragraphs[0]
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    right.add_run(f"{invoice_config['utsteder']}\n")
    right.add_run(f"{invoice_config['adr1']}\n")
    right.add_run(f"{invoice_config['adr2']}\n")
    right.add_run(f"{invoice_config['adr3']}\n")
    right.add_run(f"{invoice_config['epost']}\n")
    right.add_run(f"Tlf: {invoice_config['tlf']}\n\n")  # litt luft etter Tlf

    p_payinfo = hdr_cells[1].paragraphs[0]
    p_payinfo.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    def add_info(label, value):
        run1 = p_payinfo.add_run(f"{label}: ")
        run2 = p_payinfo.add_run(str(value) + "\n")
        run2.bold = True

    add_info("Ordrenummer", felt["Ordrenummer"])
    add_info("Ordredato", felt["Ordredato"])
    total = sum(float(dict(zip(columns, row))["Beløp"]) for row in rows)
    add_info("Beløp", format_norwegian_amount(total, felt.get("Valuta", "NOK")))
    doc.add_paragraph("")

    # Beskrivelse i bold (uten ledetekst)
    if felt['Beskrivelse']:
        p = doc.add_paragraph()
        run = p.add_run(felt['Beskrivelse'])
        run.bold = True

    doc.add_paragraph("")

    # Ordrelinjer-tabell
    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].merge(hdr[4])  # slå sammen 5 første til én bred kolonne

    hdr[0].text = "Ordrelinje"
    hdr[5].text = "Beløp"

    # Juster bredder: Ordrelinje bred, Beløp smal
    hdr[0].paragraphs[0].runs[0].bold = True
    hdr[5].paragraphs[0].runs[0].bold = True
    hdr[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for row in rows:
        data = dict(zip(columns, row))
        r = table.add_row().cells
        r[0].merge(r[4])
        r[0].text = data["Ordrelinje"]
        r[5].text = format_norwegian_amount(data["Beløp"], data.get("Valuta","NOK"))
        r[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    try:
        file_name =    f"{felt['Ordrenummer']}.docx"
        full_path = os.path.join(download_path, file_name)

        doc.save(full_path)
        show_message("Ordre lastet ned til 'Downloads'-mappen.")
    except PermissionError:
        QMessageBox.warning(
            parent,
            "Feil ved lagring",
            f"Kunne ikke lagre filen:\n{full_path}\n\n"
            "Sannsynligvis er den allerede åpen i Word.\n"
            "Lukk filen og prøv igjen."
        )

"""
def set_col_width(cell, cm_value):
    # 1 cm ≈ 567 twips
    twips = int(cm_value * 567)
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:type'), 'dxa')
    tcW.set(qn('w:w'), str(twips))
    tcPr.append(tcW)
"""

def format_norwegian_amount(amount, valuta="NOK"):
    # Bruk tusenskilletegn (mellomrom) og komma som desimal
    return f"{amount:,.2f}".replace(",", "X").replace(".", ",").replace("X", " ") + f" {valuta}"

def write_manual_invoice_pdf(parent, invoice_config, rows, columns, download_path):
    # Fellesdata fra første rad (mottaker)
    first = rows[0]
    felt = dict(zip(columns, first))

    file_name = f"{felt['Ordrenummer']}.pdf"
    full_path = os.path.join(download_path, file_name)

    # Sidestørrelse og marger
    page_width, page_height = A4
    margin_left   = 1*cm
    margin_right  = 1*cm
    margin_top    = 1.5*cm
    margin_bottom = 1.5*cm

    x_left  = margin_left
    x_right = page_width - margin_right
    y_start = page_height - margin_top

    c = canvas.Canvas(full_path, pagesize=A4)

    # Tittel
    c.setFont("Helvetica-Bold", 24)
    c.drawCentredString(page_width/2, y_start, "ORDRE")

    # Venstre: mottaker
    c.setFont("Helvetica", 10)
    y = y_start - 40
    c.drawString(x_left, y, felt["Klubb"]); y -= 12
    if felt["Adresse"]:
        c.drawString(x_left, y, felt["Adresse"]); y -= 12
    if felt["E_post"]:
        c.drawString(x_left, y, felt["E_post"]); y -= 12
    if felt["Telefonnr"]:
        c.drawString(x_left, y, f"Tlf: {felt['Telefonnr']}")

    # Høyre: utsteder + betalingsinfo
    y = y_start - 40
    c.setFont("Helvetica", 10)
    c.drawRightString(x_right, y, invoice_config["utsteder"]); y -= 12
    c.drawRightString(x_right, y, invoice_config["adr1"]); y -= 12
    c.drawRightString(x_right, y, invoice_config["adr2"]); y -= 12
    c.drawRightString(x_right, y, invoice_config["adr3"]); y -= 12
    c.drawRightString(x_right, y, invoice_config["epost"]); y -= 12
    c.drawRightString(x_right, y, f"Tlf: {invoice_config['tlf']}"); y -= 24

    # Fakturainfo med bold verdier
    def add_info(c, x_right, y, label, value, font_label="Helvetica", font_value="Helvetica-Bold", size=10):
        label_text = f"{label}: "
        value_text = str(value)

        # Widths for precise placement
        val_w = stringWidth(value_text, font_value, size)
        lab_w = stringWidth(label_text, font_label, size)

        # Draw value right-aligned at x_right
        c.setFont(font_value, size)
        c.drawString(x_right - val_w, y, value_text)

        # Draw label immediately to the left of value
        c.setFont(font_label, size)
        c.drawString(x_right - val_w - lab_w, y, label_text)

        return y - 12

    y = add_info(c, x_right, y, "Ordrenummer", felt["Ordrenummer"])
    y = add_info(c, x_right, y, "Ordredato", felt["Ordredato"])
    total = sum(float(dict(zip(columns, row))["Beløp"]) for row in rows)
    y = add_info(c, x_right, y, "Beløp", format_norwegian_amount(total, felt.get("Valuta", "NOK")))

    # Beskrivelse med linjeskift og bold
    if felt["Beskrivelse"]:
        y -= 24
        c.setFont("Helvetica-Bold", 12)
        for line in str(felt["Beskrivelse"]).splitlines():
            c.drawString(x_left, y, line)
            y -= 14

    # Ordrelinjer-tabell
    y -= 36
    c.setFont("Helvetica-Bold", 10)
    c.drawString(x_left, y, "Ordrelinje")
    c.drawCentredString(x_right-1.5*cm, y, "Beløp")

    c.setFont("Helvetica", 10)
    for row in rows:
        data = dict(zip(columns, row))
        y -= 15
        c.drawString(x_left, y, data["Ordrelinje"])
        belop = format_norwegian_amount(data["Beløp"], data.get("Valuta","NOK"))
        c.drawRightString(x_right, y, belop)

    # Totalt nederst
    y -= 30
    c.setFont("Helvetica-Bold", 12)
    total_str = format_norwegian_amount(total, felt.get("Valuta","NOK"))
    c.drawRightString(x_right, y, f"Totalt: {total_str}")

    try:
        c.save()
        show_message("Ordre lastet ned til 'Downloads'-mappen.")
    except PermissionError:
        QMessageBox.warning(
            parent,
            "Feil ved lagring",
            f"Kunne ikke lagre filen:\n{full_path}\n\n"
            "Sannsynligvis er den allerede åpen.\n"
            "Lukk filen og prøv igjen."
        )

def make_amount_per_club(parent, bundle_id, order_no_base, customer_no_base):
    logging.info("control.make_amount_per_club")
    rows, columns = sql.make_amount_per_club(parent.ctx.conn_mgr, bundle_id, order_no_base, customer_no_base)
    html = HtmlBuilder.table(rows, columns, 0, sum_columns=[2], sum_position="above")
    HtmlBuilder.download(html, "BeløpPrKlubb.html")

def make_amount_per_club_product(parent, bundle_id, order_no_base, customer_no_base):
    logging.info("control.make_amount_per_club_product")
    rows, columns = sql.make_amount_per_club_product(parent.ctx.conn_mgr, bundle_id, order_no_base, customer_no_base)
    html = HtmlBuilder.grouped_rows_in_single_table(rows, columns, 1)
    HtmlBuilder.download(html, "BeløpPrKlubbProdukt.html")

def make_amount_per_product(parent, bundle_id, order_no_base, customer_no_base):
    logging.info("control.make_amount_per_club")
    rows, columns = sql.make_amount_per_product(parent.ctx.conn_mgr, bundle_id, order_no_base, customer_no_base)
    html = HtmlBuilder.table(rows, columns, 0, [1], "above")
    HtmlBuilder.download(html, "BeløpPrProduct.html")

def make_amount_per_product_club(parent, bundle_id, order_no_base, customer_no_base):
    logging.info("control.make_amount_per_club")
    rows, columns = sql.make_amount_per_product_club(parent.ctx.conn_mgr, bundle_id, order_no_base, customer_no_base)
    html = HtmlBuilder.grouped_rows_in_single_table(rows, columns, 2)
    HtmlBuilder.download(html, "BeløpPrProduktKlubb.html")

def make_amount_per_race_product(parent, bundle_id):
    logging.info("control.make_amount_per_club")
    rows, columns = sql.make_amount_per_race_product(parent.ctx.conn_mgr, bundle_id)
    html = HtmlBuilder.grouped_rows_in_single_table(rows, columns, 0)
    HtmlBuilder.download(html, "BeløpPrLøpProdukt.html")
